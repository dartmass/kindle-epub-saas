"""
Word (.docx) → 縦書きEPUB3 変換スクリプト  v0.2
- ルビ・縦中横・縦書きCSS
- スタイル未使用の見出しをフォントサイズで推定
- 表をテキストテーブルに変換
- 画像をスキップして警告
- 空段落の適切な処理
"""
import re
import uuid
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree
import ebooklib
from ebooklib import epub


VERTICAL_CSS = """
@charset "UTF-8";

html {
  -epub-writing-mode: vertical-rl;
  writing-mode: vertical-rl;
}

body {
  font-family: "游明朝", "YuMincho", "ヒラギノ明朝 ProN", "Hiragino Mincho ProN",
               "MS 明朝", serif;
  font-size: 1em;
  line-height: 1.8;
  margin: 0;
  padding: 1em;
}

h1 { font-size: 1.5em; font-weight: bold; margin: 0 0 1em 0; }
h2 { font-size: 1.3em; font-weight: bold; margin: 0 0 0.8em 0; }
h3 { font-size: 1.1em; font-weight: bold; margin: 0 0 0.6em 0; }

p { margin: 0; text-indent: 1em; }
p.noindent { text-indent: 0; }

ruby rt { font-size: 0.5em; }

.tcy {
  -epub-text-combine: horizontal;
  text-combine-upright: all;
}

table {
  border-collapse: collapse;
  margin: 1em 0;
}
td, th {
  border: 1px solid #888;
  padding: 0.3em 0.5em;
}
th { font-weight: bold; background: #eee; }

.img-placeholder {
  border: 1px dashed #aaa;
  padding: 0.5em;
  color: #888;
  font-size: 0.8em;
}
"""

TCY_PATTERN = re.compile(r'(\d{1,2})')


def escape_html(text: str) -> str:
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;'))


def apply_tcy(text: str) -> str:
    return TCY_PATTERN.sub(r'<span class="tcy">\1</span>', text)


def extract_ruby(ruby_elem) -> str:
    base_texts, rt_texts = [], []
    for child in ruby_elem:
        tag = child.tag.split('}')[-1]
        if tag == 'rubyBase':
            base_texts += [t.text or '' for t in child.iter(qn('w:t'))]
        elif tag == 'rt':
            rt_texts += [t.text or '' for t in child.iter(qn('w:t'))]
    base = ''.join(base_texts)
    rt = ''.join(rt_texts)
    return f'<ruby>{escape_html(base)}<rt>{escape_html(rt)}</rt></ruby>' if (base and rt) else escape_html(base)


def get_run_font_size(run) -> float:
    """runのフォントサイズをptで返す。取得できなければ0"""
    try:
        sz = run._r.find(qn('w:rPr') + '/' + qn('w:sz'))
        if sz is None:
            sz = run._r.find('.//' + qn('w:sz'))
        if sz is not None:
            return int(sz.get(qn('w:val'), 0)) / 2  # half-point → pt
        if run.font.size:
            return run.font.size.pt
    except Exception:
        pass
    return 0.0


def infer_heading_level(para) -> int:
    """
    スタイル名で判定できない場合、フォントサイズと太字で見出しレベルを推定する。
    0 = 見出しではない
    """
    style_name = para.style.name if para.style else ''
    if 'Heading 1' in style_name or '見出し 1' in style_name: return 1
    if 'Heading 2' in style_name or '見出し 2' in style_name: return 2
    if 'Heading 3' in style_name or '見出し 3' in style_name: return 3
    if 'Title' in style_name or 'タイトル' in style_name: return 1

    # スタイル未使用の著者向け: フォントサイズ + 太字で推定
    if not para.runs:
        return 0
    max_size = max((get_run_font_size(r) for r in para.runs), default=0)
    is_bold = any(r.bold for r in para.runs)

    if max_size >= 18 or (max_size >= 14 and is_bold): return 1
    if max_size >= 14 or (max_size >= 12 and is_bold): return 2
    if is_bold and max_size >= 11: return 3
    return 0


def run_to_html(run) -> str:
    """通常 w:r → HTML (太字・斜体対応)"""
    text = run.text or ''
    if not text:
        return ''
    html = apply_tcy(escape_html(text))
    if run.bold:
        html = f'<strong>{html}</strong>'
    if run.italic:
        html = f'<em>{html}</em>'
    return html


def paragraph_to_html(para):
    """段落 → HTML文字列。スキップすべきなら None を返す"""
    parts = []

    for child in para._p:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            text = ''.join(t.text or '' for t in child.findall(qn('w:t')))
            if text:
                # 太字・斜体はrunオブジェクト経由で判定
                is_bold = child.find('.//' + qn('w:b')) is not None
                is_italic = child.find('.//' + qn('w:i')) is not None
                html = apply_tcy(escape_html(text))
                if is_bold:   html = f'<strong>{html}</strong>'
                if is_italic: html = f'<em>{html}</em>'
                parts.append(html)
        elif tag == 'ruby':
            parts.append(extract_ruby(child))
        elif tag == 'hyperlink':
            for r in child.findall('.//' + qn('w:r')):
                text = ''.join(t.text or '' for t in r.findall(qn('w:t')))
                if text:
                    parts.append(escape_html(text))

    content = ''.join(parts)

    # 完全に空の段落（全角スペースのみも含む）
    if not content.strip().replace('　', '').replace(' ', ''):
        return '<p>　</p>'

    level = infer_heading_level(para)
    if level > 0:
        return f'<h{level}>{content}</h{level}>'

    # インデント済みテキスト（全角スペース始まり）はnoindentクラス
    if content.startswith('　') or content.startswith(' '):
        return f'<p class="noindent">{content}</p>'

    return f'<p>{content}</p>'


def table_to_html(table) -> str:
    """docx Table → HTMLテーブル"""
    rows_html = []
    for i, row in enumerate(table.rows):
        cells_html = []
        for cell in row.cells:
            text = escape_html(cell.text.replace('\n', ' '))
            tag = 'th' if i == 0 else 'td'
            cells_html.append(f'<{tag}>{text}</{tag}>')
        rows_html.append('<tr>' + ''.join(cells_html) + '</tr>')
    return '<table>' + ''.join(rows_html) + '</table>'


def docx_to_epub(
    docx_path: str,
    output_path: str,
    title: str = None,
    author: str = None
) -> dict:
    """
    変換実行。結果として dict を返す:
      { 'ok': bool, 'warnings': [str], 'chapters': int }
    """
    warnings = []
    doc = Document(docx_path)

    # ── タイトル自動検出 ──
    if not title:
        for para in doc.paragraphs:
            if infer_heading_level(para) == 1 and para.text.strip():
                title = para.text.strip()
                break
        if not title:
            title = Path(docx_path).stem
            warnings.append(f'タイトルを自動検出できなかったため、ファイル名を使用: {title}')

    author = author or '著者不明'

    book = epub.EpubBook()
    book.set_identifier(str(uuid.uuid4()))
    book.set_title(title)
    book.set_language('ja')
    book.add_author(author)

    css = epub.EpubItem(
        uid='style_vertical',
        file_name='style/vertical.css',
        media_type='text/css',
        content=VERTICAL_CSS.encode('utf-8')
    )
    book.add_item(css)

    # ── 要素リストを構築（段落 + 表を混在処理）──
    # doc.element.body の直接の子を走査して順序を保つ
    body_children = list(doc.element.body)
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}

    html_elements = []
    image_count = 0

    for child in body_children:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p' and child in para_map:
            para = para_map[child]
            # 画像チェック
            if child.find('.//' + qn('a:blip')) is not None or \
               child.find('.//' + qn('w:drawing')) is not None:
                image_count += 1
                html_elements.append(
                    f'<p class="img-placeholder">[画像 {image_count} ※Web版では表示対応予定]</p>'
                )
                warnings.append(f'画像をスキップしました（{image_count}枚目）')
                continue
            h = paragraph_to_html(para)
            if h:
                html_elements.append(h)

        elif tag == 'tbl' and child in table_map:
            table = table_map[child]
            html_elements.append(table_to_html(table))

    # ── チャプター分割（h1 で分割）──
    chapters_data = []
    current = []
    current_title = title

    for elem in html_elements:
        if elem.startswith('<h1>'):
            if current:
                chapters_data.append((current_title, current))
            current = [elem]
            current_title = re.sub('<[^>]+>', '', elem)
        else:
            current.append(elem)

    if current:
        chapters_data.append((current_title, current))

    if not chapters_data:
        chapters_data = [(title, html_elements)]

    # ── EPUBチャプター生成 ──
    epub_chapters = []
    for i, (chap_title, elems) in enumerate(chapters_data):
        html = '\n'.join([
            '<?xml version="1.0" encoding="UTF-8"?>',
            '<!DOCTYPE html>',
            '<html xmlns="http://www.w3.org/1999/xhtml"'
            ' xmlns:epub="http://www.idpf.org/2007/ops" xml:lang="ja">',
            '<head>',
            f'<title>{escape_html(chap_title)}</title>',
            '<meta charset="UTF-8"/>',
            '<link rel="stylesheet" type="text/css" href="../style/vertical.css"/>',
            '</head>',
            '<body>',
            *elems,
            '</body>',
            '</html>',
        ])
        chap = epub.EpubHtml(
            title=chap_title,
            file_name=f'chapter_{i+1:03d}.xhtml',
            lang='ja',
            content=html.encode('utf-8')
        )
        chap.add_item(css)
        book.add_item(chap)
        epub_chapters.append(chap)

    book.toc = tuple(epub_chapters)
    book.spine = ['nav'] + epub_chapters
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(output_path, book)

    return {'ok': True, 'warnings': warnings, 'chapters': len(epub_chapters)}


if __name__ == '__main__':
    import sys
    src = sys.argv[1] if len(sys.argv) >= 2 else 'test_novel.docx'
    dst = sys.argv[2] if len(sys.argv) >= 3 else src.replace('.docx', '.epub')
    result = docx_to_epub(src, dst)
    print(f'✅ EPUB生成完了: {dst}')
    print(f'   チャプター数: {result["chapters"]}')
    if result['warnings']:
        print('⚠️  警告:')
        for w in result['warnings']:
            print(f'   - {w}')
